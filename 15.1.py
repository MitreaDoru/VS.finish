import docx
f = open('guests.txt')
print(f.readline())
doc = docx.Document('invite.docx')
y = 8
z = 0
for x in f:
    paragrafGol = doc.add_paragraph('')
    paragrafGol = doc.add_paragraph('')
    paragrafGol = doc.add_paragraph('')
    
    paragraf1 = doc.add_paragraph('It would be a plesure to have the company of '.center(100))
    doc.paragraphs[z+4].style = 'Style2'
    paragraf2 = doc.add_paragraph(x.center(115))
    doc.paragraphs[z+5].style = 'Style3'
    paragraf3 = doc.add_paragraph('at 11010 memory Lane on the Evening'.center(95))
    doc.paragraphs[z+6].style = 'Style2'
    paragraf4 = doc.add_paragraph('April 1st'.center(160))
    doc.paragraphs[z+7].style = 'Style4'
    paragraf5 = doc.add_paragraph("at 7 o'clock".center(105))
    doc.paragraphs[z+8].style = 'Style2'
    doc.paragraphs[y].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    y += 8
    z += 8
doc.save('inviteFinished.docx')
exit(0)
